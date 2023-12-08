from docx import Document
from docx.shared import Pt, RGBColor
import os


def find_and_replace_with_style(docx_path, old_string, new_string):
    """
    Find a string in a DOCX file, save its font style, replace it with another string
    while applying the saved style, and save the modified DOCX file.

    Parameters:
    - docx_path (str): Path to the input DOCX file.
    - old_string (str): String to be replaced.
    - new_string (str): Replacement string.

    Returns:
    - modified_docx_path (str): Path to the modified DOCX file.
    """
    # Load the DOCX file
    doc = Document(docx_path)

    # Iterate through paragraphs to find and replace text within runs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if old_string in run.text:
                start_index = run.text.find(old_string)
                end_index = start_index + len(old_string)

                # Save the original font attributes
                original_font = run.font

                # Clear the original text
                run.clear()

                # Insert the new string with the saved font attributes
                run.add_text(run.text[:start_index])

                new_run = paragraph.add_run(new_string)
                new_run.font.name = original_font.name
                new_run.font.size = original_font.size
                new_run.font.bold = original_font.bold
                new_run.font.italic = original_font.italic
                new_run.font.underline = original_font.underline
                new_run.font.color.rgb = original_font.color.rgb

                run.add_text(run.text[end_index:])

    # Save the modified DOCX file to a new path
    modified_docx_path = f"modified_{os.path.basename(docx_path)}"
    doc.save(modified_docx_path)

    return modified_docx_path

if __name__ == "__main__":
    # Example usage:
    input_docx_path = "A.docx"  # Replace with the path to your DOCX file
    old_string = "«30» сентября 2023 года"  # Replace with the string to be replaced
    new_string = "«91» сентября 2023 года"  # Replace with the new string

    output_docx_path = find_and_replace_with_style(input_docx_path, old_string, new_string)
    print(f"Modified DOCX file saved at: {output_docx_path}")

