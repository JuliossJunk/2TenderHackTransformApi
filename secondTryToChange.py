from docx import Document
from bson import ObjectId
from pymongo import MongoClient


def find_replace_text(doc, find_text, replace_text):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for key, value in zip(find_text, replace_text):
                if key in run.text:
                    # Replace the text
                    run.text = run.text.replace(key, value)

                    # Preserve the style of the original text
                    if run.style is not None:
                        new_run = run._element
                        new_run.clear_content()
                        new_run.text = run.text
                        for child in run._element:
                            new_run.append(child)

                        # You can add more formatting attributes here
                        if 'bold' in run.style._element.attrib:
                            new_run.style.bold = run.style.bold
                        if 'italic' in run.style._element.attrib:
                            new_run.style.italic = run.style.italic
                        # Add more formatting attributes as needed


def transform_dict(input_dict):
    print(input_dict)
    order_keys = ["Number", "Period", "Predmet", "Mesto", "IKZ", "Istochnik", "Summa", "Avans"]
    # Filter keys and values based on the order_keys
    filtered_dict = {key: input_dict[key] for key in order_keys if key in input_dict}

    return filtered_dict


def get_dict_from_mongodb(document_id):
    # Assuming you have a MongoDB collection named 'your_collection'
    # Connect to MongoDB (adjust connection string accordingly)
    client = MongoClient("mongodb://bulbaman.me:16017")
    db = client["newest_db_to_Hack"]
    collection = db["my_collection"]

    object_id = ObjectId(document_id)

    # Use the distinct method to get all unique _id values
    # document_ids = collection.distinct('_id')
    # print(document_ids)

    # Retrieve the document by _id
    document = collection.find_one({"_id": object_id})


    client.close()

    if document:
        return document
    else:
        raise ValueError(f"No document found with _id: {document_id}")


def changing_Everything(etalon_id, wanted_id):
    # Sample dictionaries
    find_dict = transform_dict(get_dict_from_mongodb(etalon_id))
    replace_dict = transform_dict(get_dict_from_mongodb(wanted_id))

    # Sample DOCX file path
    docx_file_path = 'A.docx'

    # Load the DOCX document
    doc = Document(docx_file_path)

    # Extract keys and values from dictionaries
    find_text = list(find_dict.values())
    replace_text = list(replace_dict.values())

    # Find and replace text in the document
    find_replace_text(doc, find_text, replace_text)

    # Save the modified document
    modified_docx_file_path = 'modified_A.docx'
    doc.save(modified_docx_file_path)

    print(f"Document successfully modified. Saved as: {modified_docx_file_path}")

def search_substring_in_docx(docx_path, substring):
    """
    Search for a substring in a DOCX file and return a list of occurrences.

    Parameters:
    - docx_path (str): Path to the DOCX file.
    - substring (str): Substring to search for.

    Returns:
    - list: List of occurrences (tuples) containing (paragraph_index, run_index, start_index, end_index).
    """
    occurrences = []

    # Load the DOCX document
    doc = Document(docx_path)

    for paragraph_index, paragraph in enumerate(doc.paragraphs):
        for run_index, run in enumerate(paragraph.runs):
            start_index = run.text.find(substring)
            while start_index != -1:
                end_index = start_index + len(substring)
                occurrences.append((paragraph_index, run_index, start_index, end_index))
                start_index = run.text.find(substring, end_index)

    return occurrences

# Module for debugging
def to_find(path,search):
    docx_file_path = path
    search_string = search

    result = search_substring_in_docx(docx_file_path, search_string)

    if result:
        print(f"Substring '{search_string}' found in the document at the following locations:")
        for occurrence in result:
            print(f"Paragraph {occurrence[0]}, Run {occurrence[1]}, Start: {occurrence[2]}, End: {occurrence[3]}")
    else:
        print(f"Substring '{search_string}' not found in the document.")

if __name__ == "__main__":
    # first Masterpiece, second changes
    changing_Everything("6574ddd25d85588d3f5fd5d2", "6574de485d85588d3f5fd5d6")
